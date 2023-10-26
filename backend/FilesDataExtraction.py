#Import libraries
import os
from os.path import abspath
import glob
from docx import Document
from tqdm import tqdm
import pdftotext

#External Pages Importation
import RemovePdfSpecialCharacters
import SplitPdfParagraphs
import RemoveWordSpecialChar
import CheckParagraphEnd
from cleanText import clean_text_both
import re

def extractFiles_process(process_directory):
    """
    Cette fonction prend en paramètre un ensemble des fichiers PDF, Word
    depuis la variable files qui possédera le chemin du répertoire voulu,
    extrait le nom des fichiers 
    extrait le contenu textuel
    
            """
    #On récupère le répertoire de travail 
    cur_path = os.getcwd()
    
    #Accès au dossier où il y a les documents
    os.chdir(process_directory)

    # Liste des fichiers présents dans le répertoire datasets_labellises
    files = glob.glob("*/*.*") 
    # Traitement des documents 
    print(f"Liste des fichiers Fonction : {files}")
    corpus_splits = []
    
    files_name = [] #stockage des noms de fichiers sans doublon
    nbr_fichier = 1
    #On parcours chacun des fichiers retrouvé dans files
    #On insère le 1er fichier dans le tableau
    for fichier in tqdm(files):
        #print("fichier "+str(nbr_fichier)+" : "+fichier[7:-9])
        nbr_fichier = nbr_fichier+1
        #on retire les fichiers en doublons
        skip = False
        #print(skip)
        #print(fichier,"bdd",files_name)
        for file_name in files_name :
               
            if fichier[7:-9] in file_name:
#                 if "BST2021/Etude des problématiques de modularité d’un tableau de bord de véhicules adaptable à différents types de conduite.docx" == fichier:
#                     print(fichier[8:-9])
#                     print(file_name)
                skip = True
        if skip : 
            continue
        files_name.append(fichier)

        doc_text = ""
        dict_file = dict()
        #si le doc est un PDF 
        if fichier.endswith(".pdf"): 
            with open(fichier, "rb") as f_pdf:
                doc_reader = pdftotext.PDF(f_pdf)
                all_pages = ""
                #Récuperation du texte de toutes les pages
                for page_text in doc_reader:
                    all_pages += page_text #parcours toutes les pages rassemble le text en un bloc
            #print(f"All Pages Text concatenation : {all_pages}")
            #premier nettoyage du texte
            all_pages_clean = RemovePdfSpecialCharacters.Remove_Special_Characters_In_Text(all_pages)
            #print(f"Nettoyer bloc de texte de toutes les pages  : {all_pages_clean}")

            #Une fois que le texte de toutes les pages a été récuperer nous pouvons split nos paragraphes 
            pages_split_clean = SplitPdfParagraphs.split_paragraph_pdf(all_pages_clean)
            #print(f"decomposition bloc de text ens des pages en paragraphe  : {pages_split_clean}")
            #On regroupe les paragraphes qui ont été split manuellement
            doc_text = "\n\n".join(pages_split_clean)

        #si le doc est un word (docx)
        if fichier.endswith(".docx"):
            with open(fichier, "rb") as f_word:
                doc_reader = Document(f_word)
                #récuperation du texte de toutes les pages 
                for paragraph in doc_reader.paragraphs:
                    text_clean = RemoveWordSpecialChar.clean_text_word(paragraph.text)
                    if text_clean != "":
                        #Permet de garder en un paragraphe les paragraphes où on a "...:\n..." ou "...;\n..."
                        if CheckParagraphEnd.check_end_of_paragraph(text_clean)  :
                            doc_text += text_clean + " " 
                        else: 
                            doc_text += text_clean + "\n\n"
            doc_text = "\n\n".join([para for para in doc_text.split("\n\n") if len(para) > 145])
#         #si le doc est un word (docx)
#         if fichier.endswith(".docx"):
#             with open(fichier, "rb") as f_word:
#                 doc_reader = Document(f_word)
#                 #récuperation du texte de toutes les pages 
#                 for paragraph in doc_reader.paragraphs:
#                     #Permet de garder en un paragraphe les paragraphes où on a "...:\n..." ou "...;\n..."
#                     if check_end_of_paragraph(paragraph.text)  :
#                         doc_text += paragraph.text+" " 
#                     else: doc_text += paragraph.text+"\n\n"
#             doc_text = clean_text_word(doc_text)


        #si le doc est un pptx on ne le traite pass 
        if fichier.endswith(".pptx"):   continue
#=============================================================================================================
        if fichier.endswith(".pdf") or fichier.endswith(".docx"): 
            #stockage dans le dico 
            #process_fichier = "".join(re.split(r'[/.]', fichier)[1:-1]) #pour retirer le .pdf .docx et BST20../
            #dict_file['name'] = clean_text_both(process_fichier)
            dict_file['name'] = abspath(fichier)
            dict_file['content'] = doc_text
            dict_file['document_absolute_path'] = abspath(fichier)
               # dict_file['sentences'] = split_to_sentences(str(dict_file['contenu']))
               # dict_file['embeddings'] = compute_doc_embeddings([dict_file['sentences']])
               # dict_file['embeddings'] = dict_file['embeddings'].tolist()
            corpus_splits.append(dict_file)
    #Retour au répertoire du notebook   
    os.chdir(cur_path)
    #print(f"Fin reformatage  : {corpus_splits}")
    return corpus_splits
